# -*- coding: utf-8 -*-
"""
Módulo responsável por preencher o Formulário de Controle de Mudanças (Anexo 01 POP-NO-GQ-165 Rev.13)
em conformidade com BPF e formatação padrão da ADV Farma.
"""

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ====== PADRÕES DE FORMATAÇÃO ======
FORM_FONT_NAME = "Arial"
FORM_FONT_SIZE_PT = 9
FORM_FONT_COLOR = RGBColor(89, 89, 89)  # Cinza padrão (Plano de Fundo 1 – mais escuro 35%)

HDR_FONT_NAME = "Arial"
HDR_FONT_SIZE_PT = 11
HDR_FONT_COLOR = RGBColor(0, 0, 0)  # Preto

# ==============================================================
# FUNÇÕES DE APOIO – FORMATAÇÃO E ESCRITA
# ==============================================================

def _format_paragraph(p, *, name, size_pt, color_rgb, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Aplica alinhamento e fonte ao parágrafo e seus runs."""
    p.alignment = align
    for r in (p.runs or [p.add_run("")]):
        r.font.name = name
        r.font.size = Pt(size_pt)
        r.font.color.rgb = color_rgb


def _write_cell_value(cell, value: str):
    """
    Escreve texto em uma célula da TABELA do formulário:
    - Quebra em linhas por '\n'
    - Arial 9, justificado, cor RGB(89,89,89)
    """
    value = str(value or "").strip()
    cell.text = ""
    lines = value.split("\n") if value else ["—"]

    for i, ln in enumerate(lines):
        p = cell.paragraphs[i] if i < len(cell.paragraphs) else cell.add_paragraph("")
        if not p.runs:
            p.add_run(ln)
        else:
            p.runs[0].text = ln
        _format_paragraph(p, name=FORM_FONT_NAME, size_pt=FORM_FONT_SIZE_PT, color_rgb=FORM_FONT_COLOR)


def _write_header_cm_in_place(doc: Document, numero_cm: str):
    """
    Procura no CABEÇALHO a célula cujo texto é 'CM' (ou 'CMI')
    e substitui pelo número informado (Arial 11, preto, justificado).
    """
    if not numero_cm:
        return

    target_labels = {"CM", "CMI"}
    for section in doc.sections:
        header = section.header
        for tbl in header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text = cell.text.strip().upper()
                    if text in target_labels:
                        cell.text = str(numero_cm)
                        for p in cell.paragraphs:
                            _format_paragraph(
                                p,
                                name=HDR_FONT_NAME,
                                size_pt=HDR_FONT_SIZE_PT,
                                color_rgb=HDR_FONT_COLOR,
                                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            )
                        return  # encontrado uma vez, encerra


def _format_date_ddmmyyyy(value: str) -> str:
    """Converte datas ISO (aaaa-mm-dd) para dd/mm/aaaa."""
    if not value:
        return "—"
    v = value.strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(v, fmt)
            return dt.strftime("%d/%m/%Y")
        except Exception:
            continue
    return v


# ==============================================================
# MAPAS DE CAMPOS (labels)
# ==============================================================

LABELS = {
    "DATA": "DATA",
    "SOLICITANTE": "SOLICITANTE",
    "DEPARTAMENTO": "DEPARTAMENTO",
    "TITULO": "TÍTULO MUDANÇA",
    "CARATER": "CARÁTER MUDANÇA",
    "RETORNO": "RETORNO MUDANÇA",
    "SITUACAO": "SITUAÇÃO ATUAL",
    "ALTERACAO": "ALTERAÇÃO PROPOSTA",
    "JUST_MUD": "JUSTIFICATIVA MUDANÇA",
    "DESC_ITEM": "DESCRIÇÃO ITEM",
    "NUM_CORR": "NÚMERO CORRESPONDENTE",
    "ABRANGENCIA": "ABRANGÊNCIA DA MUDANÇA",
    "REFERE": "MUDANÇA REFERE-SE Á",
    "IMPACTO": "POTENCIAL IMPACTO AVALIADO",
    "CLASSIF": "CLASSIFICAÇÃO DA CRITICIDADE",
    "JUST_CLASSIF": "JUSTIFICATIVA DA CLASSIFICAÇÃO",
    "ANEXOS": "ANEXOS",
    "PLANO": "PLANO DE IMPLEMENTAÇÃO",
    "TREIN": "EXECUÇÃO DO TREINAMENTO",
    "VOE": "VERIFICAÇÃO DE EFICÁCIA (VoE) PÓS IMPLEMENTAÇÃO",
    "RES_VOE": "RESULTADOS DA VoE",
    "OBS": "OBSERVAÇÕES FINAIS",
}

TODOS_SETORES_MODELO = [
    "Farmacêutico Responsável","Garantia da Qualidade","Operações","Diretoria ou Conselho",
    "Almoxarifado","Comercial","Compras","Controle da Qualidade","Expedição","Faturamento",
    "Financeiro / Custos","Fiscal/Contábil","Informática (TI)","Manutenção","Planejamento (PCP)",
    "Produção","Regulatório","Terceiros","Validação",
]
OBRIGATORIOS_SECAO5 = {"Farmacêutico Responsável","Diretoria ou Conselho","Regulatório"}


# ==============================================================
# FUNÇÕES PRINCIPAIS DE PREENCHIMENTO
# ==============================================================

def _set_cell_right_of_label(doc: Document, label: str, value: str) -> bool:
    """Procura a linha com o label informado e preenche a célula à direita."""
    lab_upper = label.strip().upper()
    for table in doc.tables:
        for row in table.rows:
            left = row.cells[0].text.strip().upper()
            if lab_upper == left or lab_upper in left:
                _write_cell_value(row.cells[1], value)
                return True
    return False


def _preencher_secao5(doc: Document, departamentos_pertinentes):
    """Preenche a seção 5 com 'Não aplicável' nos setores não pertinentes."""
    pertinentes = set(OBRIGATORIOS_SECAO5 | set(departamentos_pertinentes or []))
    for table in doc.tables:
        for row in table.rows:
            dept = row.cells[0].text.strip()
            if dept in TODOS_SETORES_MODELO and dept not in pertinentes:
                if len(row.cells) >= 3:
                    _write_cell_value(row.cells[1], "Não aplicável")
                    _write_cell_value(row.cells[2], "Não aplicável")


def preencher_docx_from_payload(template_path: str, payload: dict) -> bytes:
    """Função principal: preenche o template DOCX e retorna o binário do arquivo."""
    doc = Document(template_path)
    join = lambda xs: "\n".join(xs) if xs else "—"

    # ====== Cabeçalho ======
    _write_header_cm_in_place(doc, payload.get("numero_cm", "") or "")

    # ====== Seções 1 a 3 ======
    _set_cell_right_of_label(doc, LABELS["DATA"], _format_date_ddmmyyyy(payload.get("data", "")))
    _set_cell_right_of_label(doc, LABELS["SOLICITANTE"], payload.get("solicitante", "—"))
    _set_cell_right_of_label(doc, LABELS["DEPARTAMENTO"], payload.get("departamento", "—"))
    _set_cell_right_of_label(doc, LABELS["TITULO"], payload.get("titulo_mudanca", "—"))
    _set_cell_right_of_label(doc, LABELS["CARATER"], payload.get("caracter_mudanca", "Temporária"))
    _set_cell_right_of_label(doc, LABELS["RETORNO"], payload.get("retorno_mudanca_temp", "—"))
    _set_cell_right_of_label(doc, LABELS["SITUACAO"], payload.get("situacao_atual", "—"))
    _set_cell_right_of_label(doc, LABELS["ALTERACAO"], payload.get("alteracao_proposta", "—"))
    _set_cell_right_of_label(doc, LABELS["JUST_MUD"], payload.get("justificativa_mudanca", "—"))

    # ====== Seção 3 ======
    _set_cell_right_of_label(doc, LABELS["DESC_ITEM"], join(payload.get("descricoes_itens")))
    _set_cell_right_of_label(doc, LABELS["NUM_CORR"], join(payload.get("numeros_correspondentes")))
    _set_cell_right_of_label(doc, LABELS["ABRANGENCIA"], payload.get("abrangencia", "—"))
    _set_cell_right_of_label(doc, LABELS["REFERE"], join(payload.get("mudanca_refere_se")))
    _set_cell_right_of_label(doc, LABELS["IMPACTO"], join(payload.get("impactos")))
    _set_cell_right_of_label(doc, LABELS["CLASSIF"], payload.get("classificacao", "—"))
    _set_cell_right_of_label(doc, LABELS["JUST_CLASSIF"], payload.get("justificativa_classificacao", "—"))

    # ====== Seção 4 ======
    _set_cell_right_of_label(doc, LABELS["ANEXOS"], join(payload.get("anexos_aplicaveis")))

    # ====== Seção 5 ======
    _preencher_secao5(doc, payload.get("departamentos_pertinentes"))

    # ====== Seção 6 ======
    plano = payload.get("plano_implementacao") or []
    plano_numerado = "\n".join(f"{i+1}. {item}" for i, item in enumerate(plano)) if plano else "—"
    _set_cell_right_of_label(doc, LABELS["PLANO"], plano_numerado)

    if payload.get("treinamento_executado") is not None:
        _set_cell_right_of_label(
            doc, LABELS["TREIN"],
            "Sim" if payload.get("treinamento_executado") else "Não"
        )

    _set_cell_right_of_label(
        doc, LABELS["VOE"],
        f"Critérios: {payload.get('voe_criterios','—')}\nPeríodo: {payload.get('voe_periodo','—')}"
    )
    _set_cell_right_of_label(doc, LABELS["RES_VOE"], payload.get("voe_resultados_esperados", "—"))

    # ====== Seção 7 ======
    _set_cell_right_of_label(doc, LABELS["OBS"], payload.get("observacoes_finais", ""))

    # ====== Retorno ======
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
